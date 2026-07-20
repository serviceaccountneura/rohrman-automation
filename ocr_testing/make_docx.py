#!/usr/bin/env python3
"""Generate a Word (.docx) comparison report: DeepSeek-OCR vs Nemotron-OCR-v2."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches

GREEN = RGBColor(0x1E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

t = doc.add_heading("OCR Model Comparison", level=0)
sub = doc.add_paragraph()
r = sub.add_run("DeepSeek-OCR  vs  NVIDIA Nemotron-OCR-v2")
r.bold = True
r.font.size = Pt(14)
meta = doc.add_paragraph(
    "Task: extract accurate, organized information from 14 dealership financial documents "
    "(parts invoices, purchase orders, GL/journal postings, parts statements).\n"
    "Hardware: single NVIDIA Tesla T4 (15 GB).  Corpus: 14 PDFs / 24 pages."
)
meta.runs[0].italic = True
meta.runs[0].font.size = Pt(9.5)

doc.add_heading("Executive summary", level=1)
doc.add_paragraph(
    "These are two different classes of tool. DeepSeek-OCR is a 3-billion-parameter "
    "vision-language model that rebuilds document structure (tables, headings, reading order) "
    "as Markdown. Nemotron-OCR-v2 is a small (84M) classical OCR engine that detects and "
    "recognizes text regions, returning flat text plus bounding boxes and confidence."
)
for bold, rest in [
    ("Accuracy / organization: ", "DeepSeek-OCR wins on the 13 clean documents."),
    ("Latency: ", "Nemotron-OCR-v2 is ~115x faster (0.33 s vs 38.5 s per page on the T4)."),
    ("Robustness: ", "Nemotron never degenerates; DeepSeek collapsed on one dense statement."),
    ("Recommended: ", "Nemotron as fast/robust first pass + fallback; DeepSeek for structured tables on clean pages."),
]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(bold).bold = True; p.add_run(rest)


def make_table(headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        run = tbl.rows[0].cells[i].paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(10)
            if str(val).startswith("✓"): run.font.color.rgb = GREEN; run.bold = True
            elif str(val).startswith("✗"): run.font.color.rgb = RED; run.bold = True
    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return tbl


doc.add_heading("1.  What each model is", level=1)
make_table(["", "DeepSeek-OCR", "Nemotron-OCR-v2 (multilingual)"], [
    ["Type", "3B vision-language model", "84M CNN detector + recognizer + relational"],
    ["Output", "Markdown + HTML tables, reading order", "Flat text regions + bbox + confidence"],
    ["Understands structure?", "Yes - rebuilds tables / key-value", "No - pure detect + recognize"],
    ["Best at", "Layout understanding", "Fast, robust text spotting"],
], widths=[1.6, 2.4, 2.6])

doc.add_heading("2.  Latency (NVIDIA Tesla T4)", level=1)
make_table(["Metric", "DeepSeek-OCR", "Nemotron-OCR-v2"], [
    ["Avg per page", "38.5 s", "0.33 s"],
    ["Max single page", "185 s", "2.5 s"],
    ["Total (24 pages)", "~924 s (15.4 min)", "~7.8 s"],
    ["Relative speed", "1x (baseline)", "~115x faster"],
], widths=[2.0, 2.3, 2.3])
doc.add_paragraph("Note: DeepSeek-OCR ran with eager attention (no FlashAttention on Turing T4); "
                  "it would be faster on Ampere or newer.").runs[0].italic = True

doc.add_heading("3.  Accuracy", level=1)
doc.add_paragraph("Same field, Honda parts invoice, GL row 3000 (ground truth from the source PDF):")
make_table(["Source", "Account", "Amount", "Control", "Control description"], [
    ["Ground truth", "3000", "-$64.26", "H0526", "HONDA MAY28"],
    ["DeepSeek-OCR", "3000", "✓ -$64.26", "✓ H0526", "✓ HONDA MAY28 (in table)"],
    ["Nemotron-OCR-v2", "3000", "✗ -$84.28", "✗ HD528", "HONDA MAY28 (flattened)"],
], widths=[1.7, 0.9, 1.1, 1.0, 2.0])
doc.add_paragraph(
    "DeepSeek-OCR preserves row-to-column relationships; Nemotron captures the text but flattens "
    "tables into a column-wise stream, and makes more character-level errors on stylized text. "
    "Nemotron, however, captured headers/footers DeepSeek sometimes omitted."
)

doc.add_heading("4.  Pros and cons", level=1)
doc.add_heading("DeepSeek-OCR", level=2)
doc.add_paragraph("Pros", style="Intense Quote")
for s in ["Reconstructs tables with correct account-to-amount mapping.",
          "Preserves reading order, headings and key/value pairs.",
          "Output (Markdown / HTML) is directly usable as structured data.",
          "Higher character-level accuracy on stylized text."]:
    doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph("Cons", style="Intense Quote")
for s in ["Slow on a T4 (~38 s/page; up to 185 s).",
          "Catastrophic failure on one dense multi-column statement (repetition loop).",
          "Sometimes omits peripheral text (headers, footers).",
          "Heavy 3B model; needs careful memory handling."]:
    doc.add_paragraph(s, style="List Bullet")
doc.add_heading("Nemotron-OCR-v2", level=2)
doc.add_paragraph("Pros", style="Intense Quote")
for s in ["Extremely fast (~0.33 s/page; ~115x faster).",
          "Robust - never degenerated, even on the dense page that broke DeepSeek.",
          "Captures all on-page text incl. headers/footers.",
          "Per-region bounding boxes + confidence scores.",
          "Tiny model; negligible memory and latency cost."]:
    doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph("Cons", style="Intense Quote")
for s in ["No structure - tables flattened; row/column relationships lost.",
          "More character-level OCR errors on stylized/low-contrast text.",
          "Requires building a custom CUDA/C++ extension to install.",
          "You must reconstruct layout/tables yourself from coordinates."]:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("5.  Recommendation", level=1)
p = doc.add_paragraph()
p.add_run("Best practical setup: ").bold = True
p.add_run("Nemotron-OCR-v2 as a fast, robust first pass (and fallback whenever DeepSeek degenerates "
          "or times out); DeepSeek-OCR for structured table/field extraction on clean pages - ideally "
          "on an Ampere+ GPU with FlashAttention enabled.")

out = "/home/ubuntu/ocr_testing/OCR_Comparison_DeepSeek_vs_Nemotron.docx"
doc.save(out)
print("saved:", out)
