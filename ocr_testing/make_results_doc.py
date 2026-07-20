#!/opt/pytorch/bin/python3
"""
Compile results/vertex/*.json (structured output) + *.txt (clean OCR text)
into a single Word document, rendering each document's fields form-style
(like the app's Form panel), followed by its clean text.

Optionally reads results/vertex/accuracy.json to add an accuracy section.
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path("/home/ubuntu/ocr_testing/results/vertex")
OUT = Path("/home/ubuntu/ocr_testing/AP_Extraction_Results.docx")
GREY = RGBColor(0x66, 0x66, 0x66)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x1E, 0x7D, 0x32)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def label(text):
    p = doc.add_paragraph()
    r = p.add_run(text.replace("_", " ").upper())
    r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(1)
    return p


def field(name, value):
    p = doc.add_paragraph()
    r = p.add_run(f"{name.replace('_',' ').title()}: ")
    r.bold = True; r.font.size = Pt(10)
    p.add_run("" if value is None else str(value)).font.size = Pt(10)
    p.paragraph_format.space_after = Pt(1)


def render_table(rows, title):
    if not rows:
        return
    keys = []
    for r in rows:
        if isinstance(r, dict):
            for k in r:
                if k not in keys and not str(k).startswith("_"):
                    keys.append(k)
    if not keys:
        return
    label(title)
    tbl = doc.add_table(rows=1, cols=len(keys)); tbl.style = "Light Grid Accent 1"
    for i, k in enumerate(keys):
        run = tbl.rows[0].cells[i].paragraphs[0].add_run(k.replace("_", " ").title())
        run.bold = True; run.font.size = Pt(8.5)
    for r in rows:
        cells = tbl.add_row().cells
        for i, k in enumerate(keys):
            v = r.get(k) if isinstance(r, dict) else ""
            run = cells[i].paragraphs[0].add_run("" if v is None else str(v))
            run.font.size = Pt(8.5)


def render(node, prefix=""):
    """Render a structured dict form-style; lists of dicts become tables."""
    for k, v in node.items():
        if str(k).startswith("_"):
            continue
        if isinstance(v, dict):
            label(k)
            render(v, prefix + "  ")
        elif isinstance(v, list):
            if v and all(isinstance(x, dict) for x in v):
                render_table(v, k)
            else:
                field(k, ", ".join(str(x) for x in v))
        else:
            field(k, v)


# ---------------- Title ----------------
doc.add_heading("AP Automate — Extraction Results", level=0)
pm = doc.add_paragraph()
rm = pm.add_run("Model: Google Gemini 2.5 Flash (Vertex AI)")
rm.bold = True; rm.font.size = Pt(12); rm.font.color.rgb = BLUE
p = doc.add_paragraph("Two-stage Gemini 2.5 Flash pipeline (OCR -> structured JSON) on Vertex AI, "
                      "the same extraction the web app produces. One section per document: "
                      "the structured form fields, then the clean OCR text.")
p.runs[0].italic = True; p.runs[0].font.size = Pt(9.5)

# ---------------- Optional accuracy section ----------------
acc_file = SRC / "accuracy.json"
if acc_file.exists():
    acc = json.loads(acc_file.read_text())
    doc.add_heading("Accuracy summary", level=1)
    if acc.get("model"):
        pmod = doc.add_paragraph()
        rmod = pmod.add_run(acc["model"]); rmod.bold = True; rmod.font.size = Pt(10)
    doc.add_paragraph(acc.get("methodology", ""))
    overall = acc.get("overall_pct")
    if overall is not None:
        pp = doc.add_paragraph()
        r = pp.add_run(f"Projected overall production accuracy (bulk + scanned): ~{overall:.0f}%")
        r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE
    rows = acc.get("per_document", [])
    if rows:
        tbl = doc.add_table(rows=1, cols=3); tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["Document", "Accuracy", "Notes"]):
            run = tbl.rows[0].cells[i].paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9)
        for r in rows:
            c = tbl.add_row().cells
            c[0].paragraphs[0].add_run(r["doc"]).font.size = Pt(9)
            rp = c[1].paragraphs[0].add_run(f"{r['pct']:.0f}%"); rp.font.size = Pt(9); rp.bold = True
            if r["pct"] >= 100: rp.font.color.rgb = GREEN
            c[2].paragraphs[0].add_run(r.get("note", "")).font.size = Pt(8.5)
    for note in acc.get("notes", []):
        doc.add_paragraph(note, style="List Bullet")
    doc.add_page_break()

# ---------------- Per-document sections ----------------
jsons = sorted(SRC.glob("*.json"))
jsons = [j for j in jsons if j.name != "accuracy.json"]
for idx, jf in enumerate(jsons, 1):
    data = json.loads(jf.read_text())
    doc.add_heading(f"{idx}. {jf.stem}", level=1)
    dt = data.get("document_type")
    if dt:
        field("Document Type", dt)
    render(data)

    txt = (SRC / (jf.stem + ".txt"))
    if txt.exists():
        doc.add_heading("Clean OCR text", level=2)
        pre = doc.add_paragraph(txt.read_text(encoding="utf-8"))
        pre.runs[0].font.name = "Consolas"; pre.runs[0].font.size = Pt(8)
    if idx < len(jsons):
        doc.add_page_break()

doc.save(OUT)
print("saved:", OUT, "docs:", len(jsons))
